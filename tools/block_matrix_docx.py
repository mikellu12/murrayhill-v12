"""The block matrices as a Word document with real, editable tables.

Written as .docx rather than exported as an image so the numbers stay text:
they can be restyled to a journal's table format, corrected in place, and
read by a screen reader. A picture of a table is none of those things.

Layout matches the map -- rows run East 42nd at the top down to East 34th,
columns run west to east -- so the table reads as the neighbourhood does
rather than alphabetically.

    .venv/Scripts/python tools/block_matrix_docx.py
"""
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).parent
sys.path.insert(0, str(HERE.parent / "src"))
from common import RES, banner

COLUMNS = ["Madison–Park", "Park–Lexington", "Lexington–3rd", "3rd–2nd", "2nd–1st"]


def _add_matrix(doc, df, title, caption, fmt, total=None):
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    head = table.rows[0].cells
    head[0].text = "Street"
    for i, col in enumerate(COLUMNS, start=1):
        head[i].text = f"{i}\n{col}"
    for cell in head:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    for name, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(name)
        cells[0].paragraphs[0].runs[0].font.bold = True
        for i, v in enumerate(row.values, start=1):
            cells[i].text = fmt.format(v)
        for c in cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.5)

    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    if total is not None:
        doc.add_paragraph()


def main():
    banner("block matrices -> docx")
    counts = pd.read_csv(RES / "tables" / "block_matrix_counts.csv", index_col=0)
    sim = pd.read_csv(RES / "tables" / "block_matrix_sim.csv", index_col=0)

    doc = Document()
    for section in doc.sections:                 # landscape: five wide columns
        section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading("Sampling nodes and mean Street Interface Matrix by block", level=1)
    intro = doc.add_paragraph(
        "Nine cross-streets by five blocks. Rows run north to south and columns "
        "west to east, so a cell sits where its block sits. Blocks are bounded by "
        "the six avenues only; Tunnel Exit Street and Tudor City Place cross some "
        "cross-streets and not others, so including them as boundaries would leave "
        "the block set ragged and no rectangular table would exist."
    )
    for r in intro.runs:
        r.font.size = Pt(10)

    _add_matrix(
        doc, counts, "Table 1. Sampling nodes per block",
        f"Total {int(counts.values.sum())} nodes across {counts.size} blocks. "
        "Counts vary with block length rather than with anything about the "
        "streetscape.", "{:.0f}", total=True)

    doc.add_page_break()
    _add_matrix(
        doc, sim, "Table 2. Mean Street Interface Matrix per block",
        "SIM is a weighted composite of the green, morphological and permeability "
        "dimensions, each a bounded share of the 180° along-street visual field. "
        "Weights are declared rather than fitted; no measured dwell outcome exists "
        "to fit them against.", "{:.3f}")

    out = RES / "tables" / "block_matrices.docx"
    doc.save(out)
    print(f"wrote {out}")
    print(f"  Table 1: {counts.shape[0]} x {counts.shape[1]}, {int(counts.values.sum())} nodes")
    print(f"  Table 2: mean SIM {sim.values.min():.3f} to {sim.values.max():.3f}")


if __name__ == "__main__":
    main()
